# %%
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT_DIR = "data"
if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)


# --- UTILS DE ESTILO ---
def style_header_cell(cell, text, bg_color="003366"):
    cell.text = text
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg_color)
    tcPr.append(shd)
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)


# ==========================================
# 1. PÓLIZA CON RUIDO (CARÁTULA + CLAUSULAS EXTRA)
# ==========================================
def create_poliza_noisy():
    doc = Document()

    # --- PÁGINA 1: CARÁTULA (PURO RUIDO) ---
    for _ in range(3):
        doc.add_paragraph()  # Espacio

    title = doc.add_heading("SEGUROS LATAM S.A.", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("CONDICIONES GENERALES DEL SEGURO DE AUTOMÓVILES")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style.font.size = Pt(14)

    doc.add_paragraph()

    # Tabla de datos irrelevantes para la conciliación
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    data = [
        ("Registro CNSF:", "H-2299-11"),
        ("Producto:", "Cobertura Amplia Plus"),
        ("Versión:", "2026-v2"),
        ("Aviso de Privacidad:", "Consulte www.seguroslatam.com/privacidad"),
    ]
    for i, (k, v) in enumerate(data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

    # Texto legal denso (Filler)
    doc.add_paragraph()
    legal_text = doc.add_paragraph(
        "IMPORTANTE: Este documento constituye el contrato de seguro. "
        "El asegurado declara conocer el contenido del artículo 25 de la Ley sobre el Contrato de Seguro. "
        "La compañía no será responsable por siniestros ocurridos bajo los efectos del alcohol o drogas..."
    )
    legal_text.style.font.size = Pt(8)
    legal_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_page_break()  # <--- SALTO DE PÁGINA PARA SEPARAR EL RUIDO DE LA DATA

    # --- PÁGINA 2: LAS REGLAS (MEZCLADAS CON REGLAS IRRELEVANTES) ---
    doc.add_heading("SECCIÓN I: DAÑOS MATERIALES", level=1)

    # Cláusula irrelevante 1
    doc.add_heading("1.1. Rotura de Cristales", level=2)
    doc.add_paragraph(
        "Se ampara la rotura de parabrisas, aletas, medallón y quemacocos con un deducible del 20%."
    )

    # Cláusula irrelevante 2
    doc.add_heading("1.2. Robo Parcial", level=2)
    doc.add_paragraph(
        "Queda excluido el robo de espejos laterales salvo contratación de cobertura accesoria."
    )

    # --- LA CARNE: REGLAS DE REPARACIÓN (SIGNAL) ---
    doc.add_heading("1.3. Baremos de Reparación y Pintura", level=2)
    p = doc.add_paragraph(
        "Para la indemnización de daños, se aplicarán estrictamente los siguientes límites:"
    )

    # Tabla de Reglas (Ground Truth)
    table_rules = doc.add_table(rows=1, cols=3)
    table_rules.style = "Table Grid"

    headers = ["CONCEPTO", "LÍMITE MÁXIMO", "CONDICIÓN"]
    for i, h in enumerate(headers):
        style_header_cell(table_rules.rows[0].cells[i], h)

    rules = [
        ("Mano de Obra Mecánica", "$600 MXN / hora", "Según tabulador."),
        (
            "Pintura (Por Pieza)",
            "$4,000.00 MXN",
            "Tope máximo inapelable. Excedente a cargo del asegurado.",
        ),  # <--- REGLA CLAVE 1
        (
            "Partes Estructurales",
            "Según Valuación",
            "Daños ocultos requieren NOTA TÉCNICA justificativa para ser pagados.",
        ),  # <--- REGLA CLAVE 2
    ]

    for concepto, limite, condicion in rules:
        row = table_rules.add_row()
        row.cells[0].text = concepto
        row.cells[1].text = limite
        row.cells[2].text = condicion

    # Cláusula irrelevante 3
    doc.add_paragraph()
    doc.add_heading("1.4. Auto Sustituto", level=2)
    doc.add_paragraph("Limitado a 10 días naturales en caso de Pérdida Total.")

    doc.save(os.path.join(OUTPUT_DIR, "Poliza.docx"))
    print("✅ Poliza.docx generada con carátula y cláusulas de relleno.")


# ==========================================
# 2. REPORTE CON RUIDO (CHECKLIST + PUBLICIDAD)
# ==========================================
def create_reporte_noisy():
    doc = Document()

    # --- ENCABEZADO TALLER ---
    header_table = doc.add_table(rows=1, cols=2)
    header_table.width = Inches(6)
    c1 = header_table.rows[0].cells[0]
    c1.text = "SERVICIO AUTOMOTRIZ 'EL VELOZ'"
    c1.paragraphs[0].runs[0].bold = True
    c1.paragraphs[0].runs[0].font.size = Pt(16)

    c2 = header_table.rows[0].cells[1]
    c2.text = "Calle 10, Zona Industrial\nTel: 55-1234-5678\nFB: /TallerElVeloz"
    c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph("-" * 70)

    # --- RUIDO: DATOS DE RECEPCIÓN (CHECKLIST) ---
    # Esto confunde a los sistemas viejos porque ven "Radio", "Tapetes" y creen que son refacciones
    doc.add_heading("INVENTARIO DE INGRESO (NO REPARAR)", level=3)

    checklist_table = doc.add_table(rows=2, cols=4)
    checklist_table.style = "Table Grid"

    items = [
        ("Gasolina", "50%"),
        ("Radio/Frontal", "SÍ"),
        ("Tapetes", "4 Pzas"),
        ("Herramienta", "NO"),
        ("Rayones Previos", "Puerta Izq"),
        ("Antena", "SÍ"),
        ("Gato Hidráulico", "SÍ"),
        ("Extintor", "NO"),
    ]

    r = 0
    c = 0
    for label, val in items:
        cell = checklist_table.rows[r].cells[c]
        cell.text = f"{label}: {val}"
        cell.paragraphs[0].style.font.size = Pt(8)
        c += 1
        if c > 3:
            c = 0
            r += 1

    doc.add_paragraph()

    # --- SIGNAL: NARRATIVA TÉCNICA ---
    doc.add_heading("DIAGNÓSTICO DE REPARACIÓN", level=2)
    narrative = doc.add_paragraph()
    run = narrative.add_run(
        "Vehículo ingresa por colisión frontal. Se observa daño en facia. "
        "IMPORTANTE: Al desmontar la facia, se encontró el Absorbedor de Impacto fracturado verticalmente. "
        "Se anexa evidencia fotográfica en sistema. Se requiere cambio por seguridad."
    )
    run.bold = True
    # No lo ponemos rojo esta vez, dejemos que la IA lo encuentre en texto plano, es más impresionante.

    doc.add_paragraph()

    # --- SIGNAL: TABLA DE COSTOS ---
    doc.add_heading("PRESUPUESTO SOLICITADO", level=2)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    headers = ["DESCRIPCIÓN REFACCIÓN / MO", "PRECIO UNIT.", "IMPORTE"]
    for i, h in enumerate(headers):
        style_header_cell(table.rows[0].cells[i], h, "444444")  # Gris oscuro

    # ITEMS DEL DEMO (Los mismos 3 casos)
    items_factura = [
        ("Facia Delantera (OEM)", "$3,500.00", "$3,500.00"),  # OK
        ("Pintura de Facia (Bicapa)", "$5,000.00", "$5,000.00"),  # ERROR (Tope es 4000)
        ("Absorbedor de Impacto", "$1,200.00", "$1,200.00"),  # JUSTIFICADO
    ]

    for desc, precio, total in items_factura:
        row = table.add_row()
        row.cells[0].text = desc
        row.cells[1].text = precio
        row.cells[2].text = total

    # --- RUIDO FINAL: DISCLAIMERS ---
    doc.add_paragraph()
    disclaimer = doc.add_paragraph(
        "CONDICIONES DE PAGO: 50% anticipo. El taller no se hace responsable por objetos olvidados. "
        "Garantía de 30 días en mano de obra. Precios más IVA."
    )
    disclaimer.style.font.size = Pt(7)

    doc.save(os.path.join(OUTPUT_DIR, "Reporte.docx"))
    print("✅ Reporte.docx generado con checklist de inventario (ruido).")


if __name__ == "__main__":
    create_poliza_noisy()
    create_reporte_noisy()

# %%
